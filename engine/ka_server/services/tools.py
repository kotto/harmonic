"""
KA Server — Tools Service
=========================
Services utilitaire : analyse document, traduction, gestion d'idées.

Dépendances :
    - requests (stdlib)
    - os, json, uuid, datetime (stdlib)
    - harmonic_ai.ask() pour les traitements LLM
"""

import logging
import json
import os
import uuid
import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

log = logging.getLogger(__name__)

# ── Chemins ─────────────────────────────────────────────────────────────────
_DATA_DIR = Path(__file__).resolve().parent.parent.parent / 'data' / 'saas_wave'
_IDEAS_FILE = _DATA_DIR / 'ideas.json'

# ── Services hooks (set from init) ──────────────────────────────────────────
_harmonic_ai = None


def init_tools(harmonic_ai=None):
    """Initialise le service tools avec les dépendances."""
    global _harmonic_ai
    _harmonic_ai = harmonic_ai
    _DATA_DIR.mkdir(parents=True, exist_ok=True)
    log.info("  🛠️  Tools Service prêt (analyse doc, traduction, idées)")
    return True


# ═══════════════════════════════════════════════════════════════════════════════
# 1. 📄 ANALYSE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {
    '.txt': 'text/plain',
    '.md': 'text/markdown',
    '.csv': 'text/csv',
    '.json': 'application/json',
    '.xml': 'application/xml',
    '.html': 'text/html',
    '.htm': 'text/html',
    '.py': 'text/x-python',
    '.js': 'text/javascript',
    '.ts': 'text/typescript',
    '.java': 'text/x-java',
    '.c': 'text/x-c',
    '.cpp': 'text/x-c++',
    '.h': 'text/x-c-header',
    '.rb': 'text/x-ruby',
    '.go': 'text/x-go',
    '.rs': 'text/x-rust',
    '.sh': 'text/x-shellscript',
    '.yaml': 'text/yaml',
    '.yml': 'text/yaml',
    '.toml': 'text/toml',
    '.ini': 'text/x-ini',
    '.cfg': 'text/x-config',
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
}


def _extract_text(file_data: bytes, filename: str) -> tuple:
    """
    Extrait le texte d'un fichier selon son extension.
    
    Returns:
        (texte_extrait, metadata_dict)
    """
    ext = Path(filename).suffix.lower()
    metadata = {
        'filename': filename,
        'extension': ext,
        'size_bytes': len(file_data),
        'size_display': _format_size(len(file_data)),
        'mime_type': SUPPORTED_EXTENSIONS.get(ext, 'application/octet-stream'),
    }
    
    # --- Texte brut ---
    if ext in ('.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm',
               '.py', '.js', '.ts', '.java', '.c', '.cpp', '.h',
               '.rb', '.go', '.rs', '.sh', '.yaml', '.yml', '.toml',
               '.ini', '.cfg'):
        try:
            text = file_data.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = file_data.decode('latin-1')
            except Exception:
                text = file_data.decode('utf-8', errors='replace')
        metadata['encoding'] = 'utf-8'
        metadata['char_count'] = len(text)
        metadata['line_count'] = text.count('\n') + 1
        metadata['word_count'] = len(text.split())
        return text, metadata
    
    # --- PDF ---
    if ext == '.pdf':
        return _extract_pdf(file_data, metadata)
    
    # --- DOCX ---
    if ext == '.docx':
        return _extract_docx(file_data, metadata)
    
    # --- Non supporté ---
    return '', {**metadata, 'error': f'Extension {ext} non supportée pour l\'extraction de texte'}


def _extract_pdf(file_data: bytes, metadata: dict) -> tuple:
    """Extraction texte depuis PDF avec fallback."""
    text = ''
    # Essayer pypdf (léger, std)
    try:
        import io
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            pages = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text() or ''
                pages.append(t)
            text = '\n'.join(pages)
            metadata['pages'] = len(pages)
            metadata['extraction_method'] = 'PyPDF2'
        except ImportError:
            try:
                import pdfminer
                from pdfminer.high_level import extract_text as pdf_extract
                text = pdf_extract(io.BytesIO(file_data))
                metadata['extraction_method'] = 'pdfminer'
            except ImportError:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                        text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
                    metadata['pages'] = len(pdf.pages)
                    metadata['extraction_method'] = 'pdfplumber'
                except ImportError:
                    text = '[PDF] Extraction non disponible — installez PyPDF2 ou pdfminer'
                    metadata['error'] = 'No PDF library available'
    except Exception as e:
        log.warning(f"PDF extraction failed: {e}")
        text = f'[Erreur extraction PDF: {e}]'
        metadata['error'] = str(e)
    
    metadata['char_count'] = len(text)
    metadata['word_count'] = len(text.split()) if text else 0
    return text, metadata


def _extract_docx(file_data: bytes, metadata: dict) -> tuple:
    """Extraction texte depuis DOCX."""
    text = ''
    try:
        import io
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            text = '\n'.join(p.text for p in doc.paragraphs)
            metadata['extraction_method'] = 'python-docx'
            metadata['paragraphs'] = len(doc.paragraphs)
        except ImportError:
            text = '[DOCX] Extraction non disponible — installez python-docx'
            metadata['error'] = 'python-docx not available'
    except Exception as e:
        log.warning(f"DOCX extraction failed: {e}")
        text = f'[Erreur extraction DOCX: {e}]'
        metadata['error'] = str(e)
    
    metadata['char_count'] = len(text)
    metadata['word_count'] = len(text.split()) if text else 0
    return text, metadata


def _summarize_text(text: str, filename: str) -> str:
    """
    Résumé intelligent : priorité Harmonic AI, fallback extraction.
    Extraction = phrases les plus longues (souvent les plus informatives).
    """
    text = text.strip()
    if not text:
        return "[Document vide]"
    
    # Essayer Harmonic AI
    if _harmonic_ai and hasattr(_harmonic_ai, 'ask') and len(text) < 3000:
        try:
            prompt = f"Résumé du document : {text[:1500]}\n\nRésumé en 2-3 phrases :"
            result = _harmonic_ai.ask(prompt)
            if isinstance(result, dict):
                answer = result.get('answer', '')
            else:
                answer = str(result)
            if answer and len(answer) > 20 and not any(c.isdigit() for c in answer[:5]):
                return answer[:500]
        except Exception:
            pass
    
    # Fallback : extraction des phrases les plus informatives
    import re
    sentences = re.split(r'(?<=[.!?])\s+', text)
    if len(sentences) <= 3:
        return text[:500]
    
    # Trier par longueur (les phrases longues sont souvent les plus informatives)
    scored = [(len(s.strip()), s.strip()) for s in sentences if len(s.strip()) > 20]
    scored.sort(reverse=True)
    
    # Prendre les 3 meilleures phrases
    top = [s for _, s in scored[:3]]
    summary = ' '.join(top)
    
    if len(summary) > 500:
        summary = summary[:497] + '...'
    
    return summary


def analyze_document(file_data: bytes, filename: str) -> Dict[str, Any]:
    """
    Analyse et résume un document.
    
    Args:
        file_data: Contenu brut du fichier
        filename: Nom du fichier (pour détecter l'extension)
    
    Returns:
        Dict avec résumé, métadonnées, texte extrait
    """
    # 1. Extraire le texte
    text, metadata = _extract_text(file_data, filename)
    
    if not text or text.startswith('['):
        return {
            'success': False,
            'error': metadata.get('error', 'Impossible d\'extraire le texte'),
            'metadata': metadata,
        }
    
    # 2. Générer le résumé
    summary = _summarize_text(text, filename)
    
    # 3. Retourner le résultat
    return {
        'success': True,
        'summary': summary,
        'metadata': metadata,
        'text_preview': text[:500] + ('...' if len(text) > 500 else ''),
        'text_length': len(text),
        'can_analyze_deeper': len(text) > 3000,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# 2. 🌐 TRADUCTION
# ═══════════════════════════════════════════════════════════════════════════════

# Détection simple de langue par mots courants
_LANG_SIGNATURES = {
    'fr': {'bonjour', 'je', 'tu', 'il', 'elle', 'nous', 'vous', 'ils', 'le', 'la', 'les',
           'un', 'une', 'des', 'et', 'ou', 'mais', 'donc', 'car', 'ni', 'que', 'est',
           'sont', 'avec', 'pour', 'dans', 'sur', 'pas', 'plus', 'très', 'bien', 'fait'},
    'en': {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has',
           'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
           'i', 'you', 'he', 'she', 'it', 'we', 'they', 'this', 'that', 'these', 'those',
           'and', 'or', 'but', 'if', 'because', 'so', 'not', 'very', 'good', 'well'},
    'es': {'el', 'la', 'los', 'las', 'un', 'una', 'y', 'o', 'pero', 'que', 'es', 'son',
           'está', 'están', 'tiene', 'tienen', 'con', 'para', 'por', 'en', 'no', 'se',
           'le', 'lo', 'su', 'del', 'como', 'más', 'muy', 'bien', 'todo', 'cada'},
    'de': {'der', 'die', 'das', 'den', 'dem', 'des', 'ein', 'eine', 'einen', 'einer',
           'und', 'oder', 'aber', 'wenn', 'weil', 'dass', 'ist', 'sind', 'war', 'waren',
           'hat', 'haben', 'mit', 'für', 'auf', 'in', 'nicht', 'sehr', 'gut', 'auch'},
    'it': {'il', 'la', 'lo', 'gli', 'le', 'un', 'una', 'uno', 'e', 'o', 'ma', 'che',
           'è', 'sono', 'ha', 'hanno', 'con', 'per', 'in', 'non', 'si', 'più', 'molto',
           'bene', 'cosa', 'come', 'quando', 'dove', 'chi'},
    'pt': {'o', 'a', 'os', 'as', 'um', 'uma', 'e', 'ou', 'mas', 'que', 'é', 'são',
           'está', 'estão', 'tem', 'têm', 'com', 'para', 'por', 'em', 'não', 'se',
           'mais', 'muito', 'bem', 'como', 'quando', 'onde'},
}


def _detect_lang(text: str) -> str:
    """Détection simple de langue basée sur les mots les plus fréquents."""
    if not text or len(text) < 10:
        return 'unknown'
    words = set(text.lower().split())
    scores = {}
    for lang, sig in _LANG_SIGNATURES.items():
        scores[lang] = len(words & sig)
    if not scores:
        return 'unknown'
    best = max(scores, key=scores.get)
    return best if scores[best] > 2 else 'unknown'


SUPPORTED_LANGUAGES = {
    'fr': 'Français', 'en': 'English', 'es': 'Español',
    'de': 'Deutsch', 'it': 'Italiano', 'pt': 'Português',
    'nl': 'Nederlands', 'ru': 'Русский', 'zh': '中文',
    'ja': '日本語', 'ar': 'العربية', 'ko': '한국어',
}


def translate(text: str, target: str = 'en', source: Optional[str] = None) -> Dict[str, Any]:
    """
    Traduit un texte via Harmonic AI.
    
    Args:
        text: Texte à traduire
        target: Langue cible (code ISO)
        source: Langue source (auto-détection si None)
    
    Returns:
        Dict avec traduction, langues, confiance
    """
    if not text.strip():
        return {'success': False, 'error': 'Texte vide'}
    
    if target not in SUPPORTED_LANGUAGES:
        return {'success': False, 'error': f'Langue cible non supportée: {target}',
                'supported': list(SUPPORTED_LANGUAGES.keys())}
    
    # Auto-détection
    if not source:
        source = _detect_lang(text)
        if source == 'unknown':
            source = 'fr'  # fallback
    
    if source == target:
        return {
            'success': True,
            'translated_text': text,
            'source_lang': source,
            'target_lang': target,
            'note': 'Les langues source et cible sont identiques',
        }
    
    source_name = SUPPORTED_LANGUAGES.get(source, source)
    target_name = SUPPORTED_LANGUAGES.get(target, target)
    
    # Traduction : priorité API externe (MyMemory), fallback harmonic AI
    translated = ''
    api_success = False
    
    # Essayer MyMemory API (gratuit, rapide)
    try:
        import requests as req
        url = "https://api.mymemory.translated.net/get"
        params = {
            'q': text[:500],  # limite 500 caractères par requête
            'langpair': f'{source}|{target}',
        }
        resp = req.get(url, params=params, timeout=10)
        if resp.status_code == 200:
            data = resp.json()
            if data.get('responseStatus') == 200:
                translated = data['responseData']['translatedText']
                if translated:
                    api_success = True
                    log.info(f"🌐 MyMemory: {source}->{target} ({len(text)} chars)")
    except Exception as e:
        log.warning(f"MyMemory API error: {e}")
    
    # Fallback : Harmonic AI
    if not api_success and _harmonic_ai and hasattr(_harmonic_ai, 'ask'):
        try:
            prompt = (
                f"Tu es un traducteur. Traduis le texte suivant du {source_name} vers le {target_name}. "
                f"Réponds UNIQUEMENT avec la traduction, sans introduction, sans commentaire.\n\n"
                f"Texte :\n{text}"
            )
            result = _harmonic_ai.ask(prompt)
            if isinstance(result, dict):
                translated = result.get('answer', '')
            else:
                translated = str(result)
            translated = translated.strip().strip('"').strip("'").strip('«').strip('»').strip()
        except Exception as e:
            log.warning(f"Translation fallback failed: {e}")
    
    # Si rien n'a fonctionné, retourner le texte original
    if not translated:
        translated = text
    
    if not translated:
        # Fallback : retourner le texte original
        translated = text
    
    return {
        'success': bool(translated),
        'translated_text': translated,
        'source_lang': source,
        'target_lang': target,
        'source_name': source_name,
        'target_name': target_name,
        'original_text': text,
        'char_count': len(text),
        'translated_char_count': len(translated),
    }


# ═══════════════════════════════════════════════════════════════════════════════
# 3. 💡 IDÉES
# ═══════════════════════════════════════════════════════════════════════════════

def _load_ideas() -> List[Dict]:
    """Charge les idées depuis le fichier JSON."""
    if not _IDEAS_FILE.exists():
        return []
    try:
        with open(_IDEAS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        log.error(f"Erreur chargement idées: {e}")
        return []


def _save_ideas(ideas: List[Dict]):
    """Sauvegarde les idées dans le fichier JSON."""
    _DATA_DIR.mkdir(parents=True, exist_ok=True)
    tmp = str(_IDEAS_FILE) + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(ideas, f, ensure_ascii=False, indent=2)
    os.replace(tmp, str(_IDEAS_FILE))


def list_ideas(search: str = '', tag: str = '', page: int = 1, per_page: int = 20) -> Dict[str, Any]:
    """Liste les idées avec filtres et pagination."""
    ideas = _load_ideas()
    
    # Filtres
    if search:
        search_lower = search.lower()
        ideas = [i for i in ideas if search_lower in i.get('title', '').lower()
                 or search_lower in i.get('body', '').lower()]
    if tag:
        ideas = [i for i in ideas if tag in i.get('tags', [])]
    
    # Trier par date (plus récent d'abord)
    ideas.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    
    # Pagination
    total = len(ideas)
    start = (page - 1) * per_page
    end = start + per_page
    page_ideas = ideas[start:end]
    
    # Collecter tous les tags
    all_tags = list(set(tag for idea in ideas for tag in idea.get('tags', []) if tag))
    
    return {
        'success': True,
        'ideas': page_ideas,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': total,
            'pages': max(1, (total + per_page - 1) // per_page),
        },
        'tags': sorted(all_tags),
    }


def create_idea(title: str, body: str = '', tags: list = None) -> Dict[str, Any]:
    """Crée une nouvelle idée."""
    if not title or not title.strip():
        return {'success': False, 'error': 'Le titre est requis'}
    
    ideas = _load_ideas()
    now = datetime.datetime.now(datetime.timezone.utc).isoformat()
    
    idea = {
        'id': str(uuid.uuid4())[:8],
        'title': title.strip(),
        'body': body.strip() if body else '',
        'tags': [t.strip().lower() for t in (tags or []) if t.strip()],
        'created_at': now,
        'updated_at': now,
    }
    
    ideas.append(idea)
    _save_ideas(ideas)
    
    return {
        'success': True,
        'idea': idea,
        'message': f'Idée "{title}" sauvegardée (ID: {idea["id"]})',
    }


def get_idea(idea_id: str) -> Dict[str, Any]:
    """Récupère une idée par son ID."""
    ideas = _load_ideas()
    for idea in ideas:
        if idea.get('id') == idea_id:
            return {'success': True, 'idea': idea}
    return {'success': False, 'error': 'Idée non trouvée', 'code': 'NOT_FOUND'}


def update_idea(idea_id: str, title: str = None, body: str = None, tags: list = None) -> Dict[str, Any]:
    """Met à jour une idée existante."""
    ideas = _load_ideas()
    for idea in ideas:
        if idea.get('id') == idea_id:
            if title is not None:
                idea['title'] = title.strip()
            if body is not None:
                idea['body'] = body.strip()
            if tags is not None:
                idea['tags'] = [t.strip().lower() for t in tags if t.strip()]
            idea['updated_at'] = datetime.datetime.now(datetime.timezone.utc).isoformat()
            _save_ideas(ideas)
            return {'success': True, 'idea': idea, 'message': 'Idée mise à jour'}
    return {'success': False, 'error': 'Idée non trouvée', 'code': 'NOT_FOUND'}


def delete_idea(idea_id: str) -> Dict[str, Any]:
    """Supprime une idée."""
    ideas = _load_ideas()
    for i, idea in enumerate(ideas):
        if idea.get('id') == idea_id:
            deleted = ideas.pop(i)
            _save_ideas(ideas)
            return {'success': True, 'idea': deleted, 'message': f'Idée "{deleted["title"]}" supprimée'}
    return {'success': False, 'error': 'Idée non trouvée', 'code': 'NOT_FOUND'}


def _format_size(size_bytes: int) -> str:
    """Formate une taille en bytes en lisible."""
    if size_bytes < 1024:
        return f'{size_bytes} o'
    elif size_bytes < 1024 * 1024:
        return f'{size_bytes / 1024:.1f} Ko'
    else:
        return f'{size_bytes / (1024 * 1024):.1f} Mo'
