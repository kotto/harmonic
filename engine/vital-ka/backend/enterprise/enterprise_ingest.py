"""
Enterprise Ingestor — Injection de données propriétaires d'entreprise
======================================================================
Permet aux entreprises d'injecter leurs données (PDF, DOCX, CSV, JSON, TXT)
dans l'IA Harmonique pour créer une base de connaissances privée.

Architecture :
  Fichiers → Parser → TextExtractor → Bootstrapper → QualityFilter → KB(NPZ)

Usage :
  ingestor = EnterpriseIngestor(brain=brain)
  result = ingestor.ingest_files(
      file_paths=["/data/docs/contrats.pdf", "/data/docs/produits.csv"],
      enterprise_id="acme_corp",
      domain="documentation_interne"
  )
"""

import os
import re
import sys
import time
import json
import uuid
import logging
import threading
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Callable, Any
from dataclasses import dataclass, field
from datetime import datetime
from collections import Counter

import numpy as np

log = logging.getLogger(__name__)

# ── Chemins ────────────────────────────────────────────────────────────────
_ENGINE_DIR = Path(__file__).resolve().parent
_ENTERPRISES_DIR = _ENGINE_DIR / "data" / "enterprises"
_ENTERPRISES_DIR.mkdir(parents=True, exist_ok=True)
_INDEX_PATH = _ENTERPRISES_DIR / "_index.json"

# ── Imports locaux (lazy) ──────────────────────────────────────────────────
_HarmonicBrain = None
_QualityFilter = None
_BOOTSTRAPPER_LLM = None
_BOOTSTRAPPER_SIMPLE = None


def _ensure_imports():
    """Initialise les imports lazy."""
    global _HarmonicBrain, _QualityFilter, _BOOTSTRAPPER_LLM, _BOOTSTRAPPER_SIMPLE

    if _HarmonicBrain is None:
        sys.path.insert(0, str(_ENGINE_DIR))
        try:
            from harmonic_brain import HarmonicBrain
            _HarmonicBrain = HarmonicBrain
        except ImportError:
            _HarmonicBrain = None

    if _QualityFilter is None:
        try:
            from domain_specializer import QualityFilter
            _QualityFilter = QualityFilter
        except ImportError:
            _QualityFilter = None

    if _BOOTSTRAPPER_LLM is None:
        try:
            from bootstrapper import extract_triples_llm
            _BOOTSTRAPPER_LLM = extract_triples_llm
        except ImportError:
            _BOOTSTRAPPER_LLM = None

    if _BOOTSTRAPPER_SIMPLE is None:
        try:
            from bootstrapper import extract_triples_simple
            _BOOTSTRAPPER_SIMPLE = extract_triples_simple
        except ImportError:
            _BOOTSTRAPPER_SIMPLE = None


# ═══════════════════════════════════════════════════════════════════════════════
# STRUCTURES DE DONNÉES
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class IngestionJob:
    """Tâche d'ingestion entreprise en cours."""
    job_id: str
    enterprise_id: str
    domain: str
    status: str = "pending"  # parsing | extracting | filtering | encoding | done | error
    files_total: int = 0
    files_done: int = 0
    files_failed: int = 0
    triplets_extracted: int = 0
    triplets_filtered: int = 0
    started_at: str = ""
    error: Optional[str] = None

    def to_dict(self) -> Dict:
        return {
            "job_id": self.job_id,
            "enterprise_id": self.enterprise_id,
            "domain": self.domain,
            "status": self.status,
            "files_total": self.files_total,
            "files_done": self.files_done,
            "files_failed": self.files_failed,
            "triplets_extracted": self.triplets_extracted,
            "triplets_filtered": self.triplets_filtered,
            "started_at": self.started_at,
            "error": self.error,
        }


@dataclass
class IngestionResult:
    """Résultat d'une ingestion entreprise."""
    enterprise_id: str
    domain: str
    files_processed: int
    files_failed: int
    triplets_extracted: int
    triplets_after_quality: int
    kb_path: str
    quality_report: Dict = field(default_factory=dict)
    elapsed_seconds: float = 0.0
    success: bool = True
    error: Optional[str] = None
    message: str = ""

    def to_dict(self) -> Dict:
        return {
            "enterprise_id": self.enterprise_id,
            "domain": self.domain,
            "files_processed": self.files_processed,
            "files_failed": self.files_failed,
            "triplets_extracted": self.triplets_extracted,
            "triplets_after_quality": self.triplets_after_quality,
            "kb_path": self.kb_path,
            "quality": self.quality_report,
            "elapsed_seconds": round(self.elapsed_seconds, 1),
            "success": self.success,
            "error": self.error,
            "message": self.message,
        }


# ═══════════════════════════════════════════════════════════════════════════════
# PARSEURS DE DOCUMENTS
# ═══════════════════════════════════════════════════════════════════════════════

class BaseParser:
    """Classe de base pour les parseurs de documents."""

    def can_handle(self, file_path: str) -> bool:
        """Vérifie si ce parseur peut traiter le fichier."""
        return False

    def extract_text(self, file_path: str) -> Optional[str]:
        """Extrait le texte du document."""
        raise NotImplementedError

    def extract_metadata(self, file_path: str) -> Dict:
        """Extrait les métadonnées du document."""
        return {"file": file_path, "size_bytes": os.path.getsize(file_path) if os.path.exists(file_path) else 0}


class TextParser(BaseParser):
    """Parseur pour fichiers texte (.txt, .md, .rst, .log, etc.)."""

    TEXT_EXTENSIONS = {'.txt', '.md', '.rst', '.log', '.csv', '.tsv',
                       '.json', '.jsonl', '.xml', '.yaml', '.yml', '.toml',
                       '.py', '.js', '.ts', '.java', '.c', '.cpp', '.h',
                       '.html', '.css', '.sql', '.sh', '.bat', '.ps1'}

    def can_handle(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.TEXT_EXTENSIONS

    def extract_text(self, file_path: str) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            log.warning(f"TextParser: impossible de lire {file_path}: {e}")
            return None


class JSONParser(BaseParser):
    """Parseur pour fichiers JSON et JSONL."""

    def can_handle(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in {'.json', '.jsonl'}

    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extrait le contenu textuel d'un JSON.
        
        Pour JSONL : chaque ligne est un objet, on extrait les valeurs textuelles.
        Pour JSON standard : on aplatit la structure et extrait les chaînes.
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read()
            
            ext = Path(file_path).suffix.lower()
            
            if ext == '.jsonl':
                return self._parse_jsonl(raw)
            else:
                return self._parse_json(raw)
        except Exception as e:
            log.warning(f"JSONParser: impossible de lire {file_path}: {e}")
            return None

    def _parse_json(self, raw: str) -> str:
        """Parse un JSON standard et extrait les chaînes."""
        try:
            data = json.loads(raw)
            return self._extract_strings(data)
        except json.JSONDecodeError:
            return raw  # Retourner le texte brut si pas du JSON valide

    def _parse_jsonl(self, raw: str) -> str:
        """Parse un fichier JSONL (un objet JSON par ligne)."""
        texts = []
        for line in raw.split('\n'):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                texts.append(self._extract_strings(obj))
            except json.JSONDecodeError:
                texts.append(line)  # Ligne non-JSON, la garder telle quelle
        return '\n'.join(texts)

    def _extract_strings(self, obj, max_depth=3) -> str:
        """Extrait récursivement toutes les chaînes d'un objet JSON."""
        texts = []
        
        if isinstance(obj, str):
            texts.append(obj)
        elif isinstance(obj, dict) and max_depth > 0:
            for key, value in obj.items():
                if isinstance(key, str) and len(key) > 2:
                    texts.append(key)
                texts.append(self._extract_strings(value, max_depth - 1))
        elif isinstance(obj, list) and max_depth > 0:
            for item in obj[:100]:  # Limiter pour les grands tableaux
                texts.append(self._extract_strings(item, max_depth - 1))
        elif isinstance(obj, (int, float, bool)):
            pass  # Ignorer les valeurs non-textuelles
        
        return ' '.join(t for t in texts if t and len(t) > 1)


class CSVParser(BaseParser):
    """Parseur pour fichiers CSV/TSV."""

    def can_handle(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in {'.csv', '.tsv'}

    def extract_text(self, file_path: str) -> Optional[str]:
        """Convertit un CSV en texte lisible pour l'extraction de triplets."""
        try:
            import csv
            
            delimiter = '\t' if file_path.endswith('.tsv') else ','
            
            texts = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                
                if reader.fieldnames:
                    # Ajouter les en-têtes comme contexte
                    texts.append("Colonnes: " + ", ".join(reader.fieldnames))
                
                for row in reader:
                    # Convertir chaque ligne en phrase descriptive
                    parts = []
                    for key, value in row.items():
                        if value and value.strip():
                            parts.append(f"{key}: {value.strip()}")
                    if parts:
                        texts.append(" | ".join(parts))
            
            return '\n'.join(texts)
        except Exception as e:
            log.warning(f"CSVParser: impossible de lire {file_path}: {e}")
            # Fallback : lire comme texte brut
            return TextParser().extract_text(file_path)


class PDFParser(BaseParser):
    """Parseur pour fichiers PDF (nécessite PyMuPDF ou pdfplumber)."""

    def can_handle(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.pdf'

    def extract_text(self, file_path: str) -> Optional[str]:
        """Extrait le texte d'un PDF. Essaie PyMuPDF puis pdfplumber."""
        # Essayer PyMuPDF (fitz)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            texts = []
            for page in doc:
                text = page.get_text()
                if text:
                    texts.append(text)
            doc.close()
            if texts:
                log.info(f"PDFParser (PyMuPDF): {len(texts)} pages extraites de {file_path}")
                return '\n\n'.join(texts)
        except ImportError:
            pass
        except Exception as e:
            log.debug(f"PDFParser (PyMuPDF) échec: {e}")

        # Essayer pdfplumber
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
            if texts:
                log.info(f"PDFParser (pdfplumber): {len(texts)} pages extraites de {file_path}")
                return '\n\n'.join(texts)
        except ImportError:
            pass
        except Exception as e:
            log.debug(f"PDFParser (pdfplumber) échec: {e}")

        log.warning(f"PDFParser: aucune bibliothèque disponible pour {file_path}. "
                     "Installez PyMuPDF: pip install PyMuPDF")
        return None


class DOCXParser(BaseParser):
    """Parseur pour fichiers DOCX (nécessite python-docx)."""

    def can_handle(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.docx'

    def extract_text(self, file_path: str) -> Optional[str]:
        """Extrait le texte d'un DOCX."""
        try:
            from docx import Document
            doc = Document(file_path)
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Ajouter aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        texts.append(" | ".join(row_texts))
            
            if texts:
                log.info(f"DOCXParser: {len(texts)} paragraphes extraits de {file_path}")
                return '\n'.join(texts)
        except ImportError:
            log.warning(f"DOCXParser: python-docx non installé. pip install python-docx")
        except Exception as e:
            log.warning(f"DOCXParser: impossible de lire {file_path}: {e}")

        return None


# ═══════════════════════════════════════════════════════════════════════════════
# ENTERPRISE INGESTOR
# ═══════════════════════════════════════════════════════════════════════════════

class EnterpriseIngestor:
    """
    Orchestrateur d'ingestion de données d'entreprise.
    
    Supporte : TXT, MD, JSON, JSONL, CSV, TSV, PDF (optionnel), DOCX (optionnel)
    """
    
    # Extensions supportées → parser
    PARSER_MAP = {
        'pdf': PDFParser,
        'docx': DOCXParser,
        'csv': CSVParser,
        'tsv': CSVParser,
        'json': JSONParser,
        'jsonl': JSONParser,
        'txt': TextParser,
        'md': TextParser,
    }
    
    def __init__(self, brain=None):
        """
        Args:
            brain: Instance HarmonicBrain principale
        """
        _ensure_imports()
        self.brain = brain
        self._parsers: Dict[str, BaseParser] = {}
        self._jobs: Dict[str, IngestionJob] = {}
        self._lock = threading.Lock()
    
    # ── API publique ───────────────────────────────────────────────────────
    
    def ingest_files(
        self,
        file_paths: List[str],
        enterprise_id: str,
        domain: str,
        depth: str = "expert",
        on_progress: Optional[Callable] = None,
    ) -> IngestionResult:
        """
        Ingère une liste de fichiers pour une entreprise.
        
        Args:
            file_paths: Chemins des fichiers à ingérer
            enterprise_id: Identifiant de l'entreprise (ex: "acme_corp")
            domain: Domaine de connaissance (ex: "documentation_interne")
            depth: Profondeur d'extraction (influence le nombre de triplets)
            on_progress: Callback optionnel (job -> None)
        
        Returns:
            IngestionResult
        """
        t0 = time.time()
        enterprise_id = self._sanitize_id(enterprise_id)
        domain = self._sanitize_id(domain)
        
        job = IngestionJob(
            job_id=str(uuid.uuid4())[:8],
            enterprise_id=enterprise_id,
            domain=domain,
            status="parsing",
            files_total=len(file_paths),
            started_at=datetime.now().isoformat(),
        )
        
        log.info(f"[{job.job_id}] EnterpriseIngestor: {len(file_paths)} fichiers "
                 f"→ enterprise={enterprise_id}, domain={domain}")
        
        try:
            # ── Étape 1 : Parsing des fichiers ──────────────────────────
            all_texts = []
            failed_files = []
            
            for i, path in enumerate(file_paths):
                job.files_done = i + 1
                job.status = "parsing"
                if on_progress:
                    on_progress(job)
                
                try:
                    text = self._parse_file(path)
                    if text and len(text.strip()) > 50:
                        all_texts.append(text)
                        log.debug(f"  ✅ {Path(path).name}: {len(text)} caractères")
                    else:
                        failed_files.append(path)
                        job.files_failed += 1
                        log.debug(f"  ⚠️ {Path(path).name}: texte vide ou trop court")
                except Exception as e:
                    failed_files.append(path)
                    job.files_failed += 1
                    log.warning(f"  ❌ {Path(path).name}: {e}")
            
            if not all_texts:
                return IngestionResult(
                    enterprise_id=enterprise_id, domain=domain,
                    files_processed=len(file_paths) - len(failed_files),
                    files_failed=len(failed_files),
                    triplets_extracted=0, triplets_after_quality=0,
                    kb_path="", success=False,
                    error="Aucun texte extrait des fichiers",
                    message="❌ Aucun contenu exploitable trouvé dans les fichiers.",
                )
            
            combined_text = '\n\n'.join(all_texts)
            log.info(f"[{job.job_id}] Parsing: {len(all_texts)} fichiers → "
                     f"{len(combined_text)} caractères, {len(failed_files)} échecs")
            
            # ── Étape 2 : Extraction de triplets ────────────────────────
            job.status = "extracting"
            if on_progress:
                on_progress(job)
            
            triplets = self._extract_triplets(combined_text, depth)
            job.triplets_extracted = len(triplets)
            log.info(f"[{job.job_id}] Extraction: {len(triplets)} triplets")
            
            if len(triplets) < 10:
                return IngestionResult(
                    enterprise_id=enterprise_id, domain=domain,
                    files_processed=len(file_paths) - len(failed_files),
                    files_failed=len(failed_files),
                    triplets_extracted=len(triplets), triplets_after_quality=0,
                    kb_path="", success=False,
                    error=f"Seulement {len(triplets)} triplets extraits",
                    message=f"⚠️ Peu de connaissances extraites ({len(triplets)} triplets). "
                            "Vérifiez le contenu des fichiers.",
                )
            
            # ── Étape 3 : Filtrage qualité ──────────────────────────────
            job.status = "filtering"
            if on_progress:
                on_progress(job)
            
            qf = _QualityFilter(enable_coherence=(len(triplets) <= 10000)) if _QualityFilter else None
            
            if qf:
                # Construire des sources fictives (une par fichier)
                sources = [{"url": f"enterprise://{enterprise_id}/{Path(p).name}",
                           "source_type": "enterprise",
                           "title": Path(p).name}
                          for p in file_paths]
                
                filtered = qf.filter(triplets, sources, domain)
                scored_triplets = filtered.triplets
                quality_report = filtered.to_dict()
                job.triplets_filtered = len(scored_triplets)
                log.info(f"[{job.job_id}] QualityFilter: {len(triplets)} → "
                         f"{len(scored_triplets)} (G={filtered.gold_count} "
                         f"S={filtered.silver_count} B={filtered.bronze_count})")
            else:
                scored_triplets = [(s, r, o, sec, 1.0) for s, r, o, sec in triplets]
                quality_report = {}
                job.triplets_filtered = len(scored_triplets)
            
            # ── Étape 4 : Construction KB ───────────────────────────────
            job.status = "encoding"
            if on_progress:
                on_progress(job)
            
            kb_path = self._build_enterprise_kb(
                enterprise_id, domain, scored_triplets, quality_report
            )
            
            # ── Résultat ────────────────────────────────────────────────
            elapsed = time.time() - t0
            job.status = "done"
            
            message = (
                f"✅ **Ingestion terminée : {domain}**\n\n"
                f"📁 Fichiers traités : {len(file_paths) - len(failed_files)}/"
                f"{len(file_paths)}\n"
                f"🧬 Triplets extraits : {len(triplets):,}\n"
                f"🛡️ Après qualité : {len(scored_triplets):,} "
                f"(⭐{quality_report.get('gold',0)} 🔶{quality_report.get('silver',0)} "
                f"🔹{quality_report.get('bronze',0)})\n"
                f"⏱️ Temps : {elapsed:.0f}s\n\n"
                f"💾 Base sauvegardée : `{kb_path}`"
            )
            
            result = IngestionResult(
                enterprise_id=enterprise_id, domain=domain,
                files_processed=len(file_paths) - len(failed_files),
                files_failed=len(failed_files),
                triplets_extracted=len(triplets),
                triplets_after_quality=len(scored_triplets),
                kb_path=str(kb_path),
                quality_report=quality_report,
                elapsed_seconds=elapsed,
                success=True,
                message=message,
            )
            
            # Charger dans le brain si disponible
            if self.brain is not None and _HarmonicBrain is not None:
                try:
                    self.brain.load_user_kb(enterprise_id, str(kb_path))
                    log.info(f"[{job.job_id}] KB chargée dans le brain: {enterprise_id}")
                except Exception as e:
                    log.warning(f"[{job.job_id}] Échec chargement brain: {e}")
            
            return result
            
        except Exception as e:
            log.exception(f"[{job.job_id}] Erreur ingestion: {e}")
            job.status = "error"
            job.error = str(e)
            return IngestionResult(
                enterprise_id=enterprise_id, domain=domain,
                files_processed=0, files_failed=len(file_paths),
                triplets_extracted=0, triplets_after_quality=0,
                kb_path="", success=False, error=str(e),
                message=f"❌ Erreur : {e}",
            )
    
    def ingest_directory(
        self,
        dir_path: str,
        enterprise_id: str,
        domain: str,
        depth: str = "expert",
    ) -> IngestionResult:
        """
        Ingère tous les fichiers d'un répertoire.
        
        Args:
            dir_path: Chemin du répertoire
            enterprise_id: Identifiant entreprise
            domain: Domaine de connaissance
            depth: Profondeur d'extraction
        """
        dir_path = Path(dir_path)
        if not dir_path.exists() or not dir_path.is_dir():
            return IngestionResult(
                enterprise_id=enterprise_id, domain=domain,
                files_processed=0, files_failed=0,
                triplets_extracted=0, triplets_after_quality=0,
                kb_path="", success=False,
                error=f"Répertoire introuvable: {dir_path}",
            )
        
        # Trouver tous les fichiers supportés
        supported_exts = set(self.PARSER_MAP.keys())
        file_paths = []
        for ext in supported_exts:
            file_paths.extend(str(p) for p in dir_path.glob(f"*.{ext}"))
        
        # Trier par taille (les plus gros d'abord)
        file_paths.sort(key=lambda p: os.path.getsize(p) if os.path.exists(p) else 0, reverse=True)
        
        log.info(f"EnterpriseIngestor.directory: {len(file_paths)} fichiers trouvés dans {dir_path}")
        return self.ingest_files(file_paths, enterprise_id, domain, depth)
    
    # ── Statut ──────────────────────────────────────────────────────────
    
    def get_enterprise_status(self, enterprise_id: str) -> Dict:
        """Retourne le statut d'une entreprise."""
        enterprise_id = self._sanitize_id(enterprise_id)
        ent_dir = _ENTERPRISES_DIR / enterprise_id
        profile_path = ent_dir / "enterprise_profile.json"
        
        if not profile_path.exists():
            return {
                "enterprise_id": enterprise_id,
                "exists": False,
                "domains": {},
                "total_triplets": 0,
            }
        
        try:
            with open(profile_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return {
                "enterprise_id": enterprise_id,
                "exists": True,
                "domains": data.get("domains", {}),
                "total_triplets": data.get("total_triplets", 0),
                "last_ingestion": data.get("last_ingestion", ""),
            }
        except Exception:
            return {"enterprise_id": enterprise_id, "exists": False}
    
    def list_enterprises(self) -> List[Dict]:
        """Liste toutes les entreprises avec leurs stats."""
        enterprises = []
        
        if _INDEX_PATH.exists():
            try:
                with open(_INDEX_PATH, 'r', encoding='utf-8') as f:
                    index = json.load(f)
                for ent_id in index.get("enterprises", []):
                    status = self.get_enterprise_status(ent_id)
                    if status["exists"]:
                        enterprises.append(status)
            except Exception:
                pass
        
        # Aussi scanner le dossier
        if _ENTERPRISES_DIR.exists():
            for d in _ENTERPRISES_DIR.iterdir():
                if d.is_dir() and not d.name.startswith('.') and not d.name.startswith('_'):
                    ent_id = d.name
                    if not any(e["enterprise_id"] == ent_id for e in enterprises):
                        status = self.get_enterprise_status(ent_id)
                        if status["exists"]:
                            enterprises.append(status)
        
        return enterprises
    
    # ── Internes ────────────────────────────────────────────────────────
    
    def _parse_file(self, file_path: str) -> Optional[str]:
        """Parse un fichier avec le parseur approprié."""
        ext = Path(file_path).suffix.lower().lstrip('.')
        
        # Obtenir ou créer le parseur
        if ext not in self._parsers:
            parser_cls = self.PARSER_MAP.get(ext, TextParser)
            self._parsers[ext] = parser_cls()
        
        parser = self._parsers[ext]
        return parser.extract_text(file_path)
    
    def _extract_triplets(
        self,
        text: str,
        depth: str = "expert",
    ) -> List[Tuple[str, str, str, str]]:
        """
        Extrait les triplets d'un texte long.
        
        Découpe en segments, extrait de chaque segment, déduplique.
        """
        # Découper le texte en segments (max 3000 caractères par segment)
        max_segment_len = 3000
        segments = []
        
        paragraphs = text.split('\n\n')
        current = ""
        for para in paragraphs:
            if len(current) + len(para) < max_segment_len:
                current += para + '\n\n'
            else:
                if current:
                    segments.append(current)
                current = para + '\n\n'
        if current:
            segments.append(current)
        
        # Limiter le nombre de segments selon la profondeur
        depth_limits = {"debutant": 20, "avance": 50, "expert": 100, "encyclopedique": 200}
        max_segments = depth_limits.get(depth, 100)
        segments = segments[:max_segments]
        
        all_triplets = []
        seen = set()
        
        extract_fn = _BOOTSTRAPPER_LLM if _BOOTSTRAPPER_LLM else _BOOTSTRAPPER_SIMPLE
        
        for i, segment in enumerate(segments):
            if not segment or len(segment.strip()) < 50:
                continue
            
            try:
                if extract_fn:
                    triples = extract_fn(segment)
                else:
                    triples = self._basic_extract(segment)
                
                for s, r, o, sec in triples:
                    s_clean = s.strip().lower()
                    r_clean = r.strip().lower()
                    o_clean = o.strip().lower()
                    
                    if not s_clean or not r_clean or not o_clean:
                        continue
                    if len(s_clean) < 2 or len(o_clean) < 2:
                        continue
                    if s_clean == o_clean:
                        continue
                    
                    key = (s_clean, r_clean, o_clean)
                    if key not in seen:
                        seen.add(key)
                        all_triplets.append((s_clean, r_clean, o_clean, sec))
                
                if (i + 1) % 10 == 0:
                    log.debug(f"  Extraction: {len(all_triplets)} triplets ({i+1}/{len(segments)} segments)")
                    
            except Exception as e:
                log.debug(f"  Erreur extraction segment {i+1}: {e}")
                continue
        
        log.info(f"_extract_triplets: {len(all_triplets)} triplets de {len(segments)} segments")
        return all_triplets
    
    def _basic_extract(self, text: str) -> List[Tuple[str, str, str, str]]:
        """Extraction basique par patterns regex (fallback)."""
        triples = []
        text_clean = re.sub(r'\([^)]*\)', '', text)
        text_clean = re.sub(r'\[[^\]]*\]', '', text_clean)
        
        patterns = [
            (r'([A-ZÀ-Ü][a-zà-ü]{2,30})\s+est\s+(?:un|une)\s+([a-zà-ü\s]{3,60})',
             "est un", "classification"),
            (r'([A-ZÀ-Ü][a-zà-ü]{2,30})\s+a\s+(?:découvert|inventé|créé|fondé|développé)\s+'
             r'(?:le |la |les |l\')?([a-zà-ü\s]{3,60})',
             "a créé", "découverte"),
            (r'([A-Za-zà-ü]{3,40})\s+permet\s+(?:de|d\')\s+([a-zà-ü\s]{3,60})',
             "permet de", "fonction"),
            (r'([A-Za-zà-ü]{3,40})\s+(?:est|sont)\s+composé(?:e?s)?\s+(?:de|d\')\s+'
             r'([a-zà-ü\s]{3,60})',
             "est composé de", "composition"),
        ]
        
        for pattern, relation, secteur in patterns:
            for match in re.finditer(pattern, text_clean, re.IGNORECASE):
                sujet = match.group(1).strip().lower()
                objet = match.group(2).strip().lower()
                if len(sujet) >= 2 and len(objet) >= 2:
                    triples.append((sujet, relation, objet, secteur))
        
        return triples
    
    def _build_enterprise_kb(
        self,
        enterprise_id: str,
        domain: str,
        scored_triplets: List[Tuple],
        quality_report: Dict,
    ) -> Path:
        """Construit et sauvegarde la KB entreprise."""
        ent_dir = _ENTERPRISES_DIR / enterprise_id
        ent_dir.mkdir(parents=True, exist_ok=True)
        
        # Extraire les triplets sans amplitude pour le brain
        facts_for_brain = [(s, r, o, sec) for s, r, o, sec, _amp in scored_triplets]
        
        # Sauvegarder en NPZ
        kb_filename = f"kb_{domain}.npz"
        kb_path = ent_dir / kb_filename
        facts_array = np.array(facts_for_brain, dtype=object)
        np.savez(str(kb_path), facts=facts_array)
        
        # Sauvegarder le rapport de qualité
        quality_path = ent_dir / f"quality_{domain}.json"
        with open(quality_path, 'w', encoding='utf-8') as f:
            json.dump(quality_report, f, indent=2, ensure_ascii=False)
        
        # Mettre à jour le profil entreprise
        profile_path = ent_dir / "enterprise_profile.json"
        profile_data = {}
        if profile_path.exists():
            try:
                with open(profile_path, 'r', encoding='utf-8') as f:
                    profile_data = json.load(f)
            except Exception:
                profile_data = {}
        
        if "domains" not in profile_data:
            profile_data["domains"] = {}
        
        total_triplets = sum(
            d.get("triplets", 0) for d in profile_data["domains"].values()
        )
        
        profile_data["domains"][domain] = {
            "kb_path": str(kb_path),
            "triplets": len(scored_triplets),
            "quality": {
                "gold": quality_report.get("gold", 0),
                "silver": quality_report.get("silver", 0),
                "bronze": quality_report.get("bronze", 0),
            },
            "ingested_at": datetime.now().isoformat(),
        }
        profile_data["total_triplets"] = total_triplets + len(scored_triplets)
        profile_data["last_ingestion"] = datetime.now().isoformat()
        profile_data["enterprise_id"] = enterprise_id
        
        with open(profile_path, 'w', encoding='utf-8') as f:
            json.dump(profile_data, f, indent=2, ensure_ascii=False)
        
        # Mettre à jour l'index global
        self._update_index(enterprise_id)
        
        size_mb = kb_path.stat().st_size / (1024 * 1024)
        log.info(f"KB entreprise sauvegardée: {kb_path} ({size_mb:.1f} MB, "
                 f"{len(scored_triplets)} triplets)")
        
        return kb_path
    
    def _update_index(self, enterprise_id: str):
        """Met à jour l'index global des entreprises."""
        index = {}
        if _INDEX_PATH.exists():
            try:
                with open(_INDEX_PATH, 'r', encoding='utf-8') as f:
                    index = json.load(f)
            except Exception:
                index = {}
        
        if "enterprises" not in index:
            index["enterprises"] = []
        
        if enterprise_id not in index["enterprises"]:
            index["enterprises"].append(enterprise_id)
            index["updated_at"] = datetime.now().isoformat()
        
        with open(_INDEX_PATH, 'w', encoding='utf-8') as f:
            json.dump(index, f, indent=2, ensure_ascii=False)
    
    @staticmethod
    def _sanitize_id(name: str) -> str:
        """Nettoie un identifiant (enterprise_id ou domain)."""
        return re.sub(r'[^a-z0-9_-]', '_', name.lower().strip())[:64]


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN (test)
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
    
    print("═══ Test EnterpriseIngestor ═══")
    print()
    
    # Test parsers
    import tempfile
    
    with tempfile.TemporaryDirectory() as tmp:
        # Créer un fichier test
        test_file = Path(tmp) / "test.txt"
        test_file.write_text(
            "Paris est la capitale de la France.\n"
            "L'entreprise a été fondée en 2020 par Jean Dupont.\n"
            "Le produit phare est la plateforme Harmonic AI.\n"
            "L'équipe est composée de 50 ingénieurs et 10 designers.\n"
        )
        
        ingestor = EnterpriseIngestor()
        
        # Test parsing
        text = ingestor._parse_file(str(test_file))
        print(f"  Parsing TXT: {len(text) if text else 0} caractères")
        
        # Test extraction
        triples = ingestor._extract_triplets(text, "debutant")
        print(f"  Triplets extraits: {len(triples)}")
        for s, r, o, sec in triples[:5]:
            print(f"    • {s} | {r} | {o} ({sec})")
    
    print()
    print("✅ Tests de base OK")
