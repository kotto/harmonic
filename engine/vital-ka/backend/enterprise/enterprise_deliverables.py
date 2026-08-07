#!/usr/bin/env python3
"""
enterprise_deliverables.py — Données privées → livrables (KA Enterprise)
=========================================================================

L'IA de l'entreprise répond à TOUTE question sur ses données privées et les
formate en livrables :

  - TABLEAUX : listes d'enregistrements (clients, fournisseurs, factures…),
    comptages, totaux, moyennes, minimums, maximums → Excel (.xlsx) ou CSV
    téléchargeables, avec feuille Résumé (KPIs).
  - TEXTES : rapports, emails, comptes-rendus, lettres, notes internes
    rédigés en français propre à partir du savoir du département.

Tout est déterministe (0 LLM, 0 GPU) : rappel ondulatoire (motifs + ψ),
connexion par phase (HWAT) et correction française (polish_prose).

Usage (endpoints /api/enterprise/departments/<id>/…) :
  - query_data       → aperçu JSON (colonnes, lignes, agrégats)
  - build_excel      → fichier .xlsx (fallback CSV)
  - compose_document → texte structuré (email, rapport, CR, lettre, note)
  - document_to_docx → fichier .docx (fallback .txt)
  - summarize_department → synthèse du savoir du département
"""

import io
import re
import csv
import time
from collections import Counter
from datetime import date
from typing import Dict, List, Optional, Tuple

# ═══════════════════════════════════════════════════════════════════════════════
# INTENTIONS DE DONNÉES DÉTECTÉES DANS LA QUESTION
# ═══════════════════════════════════════════════════════════════════════════════

_INTENT_KEYWORDS: Dict[str, List[str]] = {
    'compte': ['combien', 'nombre', 'comptez', 'effectif', 'volume', 'combien de lignes'],
    'somme': ['chiffre d affaires', 'somme', 'cumul', 'total', 'montant total',
              'additionner', 'addition', 'cumuler'],
    'moyenne': ['moyenne', 'moyen'],
    'max': ['maximum', 'plus eleve', 'plus grande', 'plus grand', 'meilleur',
            'le plus haut', 'record'],
    'min': ['minimum', 'plus bas', 'plus petite', 'plus petit', 'le plus faible'],
}

# Colonnes numériques courantes (pour choisir la colonne cible d'un agrégat)
_NUM_COLUMN_HINTS = ['ca', 'montant', 'chiffre', 'prix', 'cout', 'coût', 'salaire',
                     'valeur', 'quantite', 'quantité', 'stock', 'nombre', 'total',
                     'moyenne', 'revenu', 'budget', 'effectif', 'note', 'age', 'âge',
                     'duree', 'durée', 'poids', 'surface', 'volume', 'annee', 'année']

_STOPWORDS = {
    'les', 'des', 'une', 'dans', 'avec', 'pour', 'cette', 'votre', 'notre',
    'nous', 'vous', 'sont', 'tous', 'toutes', 'tout', 'quelle', 'quelles',
    'quel', 'quels', 'donner', 'donne', 'donnes', 'avoir', 'faire', 'fait',
    'liste', 'listes', 'infos', 'information', 'informations', 'detail',
    'details', 'donnees', 'concernant', 'sur', 'au', 'aux', 'est', 'et',
}

_FORMATS = ['email', 'rapport', 'compte_rendu', 'lettre', 'note']


# ═══════════════════════════════════════════════════════════════════════════════
# RAPPEL COMPLET — tous les faits pertinents (pas plafonné à k)
# ═══════════════════════════════════════════════════════════════════════════════

def _normalize(text: str) -> str:
    """Minuscules + suppression des accents (comparaisons d'intention)."""
    s = text.lower()
    for a, b in [('é', 'e'), ('è', 'e'), ('ê', 'e'), ('ë', 'e'), ('à', 'a'),
                 ('â', 'a'), ('î', 'i'), ('ï', 'i'), ('ô', 'o'), ('ù', 'u'),
                 ('û', 'u'), ('ç', 'c')]:
        s = s.replace(a, b)
    return s


def _stem(w: str) -> str:
    """Stemming minimal français : « clients » → « client », « client_1 » → « client »."""
    w = w.rstrip('0123456789_')
    if len(w) > 4 and w.endswith('s') and not w.endswith('ss'):
        w = w[:-1]
    return w


_ACCENT_MAP = str.maketrans(
    {'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e', 'à': 'a', 'â': 'a', 'î': 'i',
     'ï': 'i', 'ô': 'o', 'ù': 'u', 'û': 'u', 'ç': 'c', 'œ': 'oe', 'æ': 'ae'})


def _wordset(text: str) -> set:
    """Ensemble des mots significatifs (déaccentués + stemmés) d'un texte."""
    out = set()
    for w in re.sub(r'[^\w\s]', ' ', text.lower()).split():
        s = _stem(w.translate(_ACCENT_MAP))
        if len(s) > 2:
            out.add(s)
    return out


def _question_words(question: str) -> set:
    return _wordset(question) - _STOPWORDS


def detect_intent(question: str) -> Optional[str]:
    """Détecte l'intention de données : compte, somme, moyenne, max, min."""
    q = _normalize(question)
    for intent in ('compte', 'somme', 'moyenne', 'max', 'min'):
        for kw in _INTENT_KEYWORDS[intent]:
            if kw in q:
                return intent
    return None


def _recall_all(engine, department_id: str, question: str,
                max_facts: int = 300, min_score: float = 0.02) -> List:
    """
    Rappel complet : TOUS les faits pertinents (résonance ψ + overlap de
    mots-clés, formule du moteur), triés par score — pour les questions
    « liste de… » qui exigent la totalité des enregistrements.
    """
    import numpy as np
    dept_facts = engine.facts.get(department_id, [])
    if not dept_facts:
        return []

    dept = engine.departments.get(department_id)
    psi_q = engine._text_to_psi(question)
    if dept:
        psi_q = psi_q * np.exp(1j * dept.phase_offset)
    q_norm = psi_q / (np.linalg.norm(psi_q) + 1e-10)

    q_words = _question_words(question)
    # Pour une agrégation (« chiffre d'affaires total »), élargir avec le nom
    # du département : on veut TOUTES les lignes, pas seulement celles qui
    # contiennent les mots de la question.
    if detect_intent(question) and dept:
        q_words |= _question_words(dept.name)

    if not q_words:
        return dept_facts[:max_facts]

    scored = []
    for i, fact in enumerate(dept_facts):
        f_norm = fact.psi_vector / (np.linalg.norm(fact.psi_vector) + 1e-10)
        psi_score = np.real(np.dot(q_norm, np.conj(f_norm)))
        f_words = _wordset(fact.text)
        overlap = len(q_words & f_words) / max(len(q_words), 1)
        combined = 0.7 * overlap + 0.3 * max(0, psi_score)
        scored.append((combined, overlap, i, fact))

    # « matched » = au moins un mot commun (le ψ seul ne suffit jamais à
    # prouver la pertinence d'une liste) ; tri par overlap décroissant, puis
    # par ordre d'ingestion (une liste doit rester dans l'ordre d'origine).
    matched = [(ov, i, f) for s, ov, i, f in scored if s >= min_score and ov > 0]
    rest = [(ov, i, f) for s, ov, i, f in scored if not (s >= min_score and ov > 0)]
    matched.sort(key=lambda x: (-x[0], x[1]))
    if detect_intent(question):
        # Agrégation (compte, total, moyenne…) : toutes les lignes du lot,
        # les correspondantes d'abord — les autres lignes du département
        # restent dans l'ordre d'ingestion.
        picked = [f for _, _, f in matched] + [f for _, _, f in rest]
    else:
        # Liste : uniquement les faits correspondants ; fallback sur le
        # département entier si aucune correspondance (introuvable sinon).
        picked = [f for _, _, f in matched] or [f for _, _, f in rest]
    return picked[:max_facts]


# ═══════════════════════════════════════════════════════════════════════════════
# ANALYSE DE STRUCTURE — les faits deviennent des lignes de tableau
# ═══════════════════════════════════════════════════════════════════════════════

_SEG_RE = re.compile(r'\s*\|\s*')
_KV_RE = re.compile(r'^\s*(?P<k>[^\s|:][^|:]{0,50}?)\s*[:：]\s*(?P<v>.+?)\s*$')


def _split_segments(text: str) -> List[str]:
    return [s.strip() for s in _SEG_RE.split(text) if s and s.strip()]


def _parse_kv(segment: str) -> Optional[Tuple[str, str]]:
    m = _KV_RE.match(segment)
    if m and m.group('v').strip() and len(m.group('k')) <= 40:
        return m.group('k').strip(), m.group('v').strip()
    return None


def _to_number(value) -> Optional[float]:
    """Parse un nombre au format FR (« 1 234,56 ») ou EN (« 1234.56 »)."""
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    s = s.replace('\u202f', '').replace('\xa0', '').replace(' ', '')
    if ',' in s and '.' in s:
        s = s.replace('.', '').replace(',', '.')   # 1.234,56
    elif ',' in s:
        s = s.replace(',', '.')
    try:
        return float(s)
    except ValueError:
        return None


def _classify_fact(text: str) -> Tuple[str, object]:
    """Classe un fait : (mode, ligne) — clefs / positionnel / texte."""
    segs = _split_segments(text)
    if len(segs) >= 2:
        kvs = [_parse_kv(s) for s in segs]
        if all(kv is not None for kv in kvs):
            return 'clefs', dict(kvs)
        if not any(kv is not None for kv in kvs):
            return 'positionnel', segs
    return 'texte', [text]


def _column_name(col: str) -> str:
    """Nom de colonne normalisé pour les comparaisons."""
    return _normalize(col).replace(' ', '')


def analyze_rows(facts: List, question: str) -> Dict:
    """
    Transforme les faits rappelés en tableau : mode (clefs / positionnel /
    texte), colonnes (union ordonnée), lignes (dicts par colonne).
    """
    modes = []
    raw = []
    for f in facts:
        mode, row = _classify_fact(f.text)
        modes.append(mode)
        raw.append(row)

    # Mode majoritaire du lot
    counts = Counter(modes)
    mode = counts.most_common(1)[0][0]

    columns: List[str] = []
    rows: List[Dict] = []
    for i, (m, row) in enumerate(zip(modes, raw)):
        if m == 'clefs':
            for k in row:
                if k not in columns:
                    columns.append(k)
        elif m == 'positionnel':
            for j in range(len(row)):
                name = f'Colonne {j + 1}'
                if name not in columns:
                    columns.append(name)

    if not columns:
        columns = ['Contenu']

    # Normalisation : toutes les lignes en dicts {colonne: valeur}
    for m, row in zip(modes, raw):
        d = {}
        if m == 'clefs':
            d = {k: row.get(k, '') for k in columns}
        elif m == 'positionnel':
            d = {columns[j] if j < len(columns) else f'Colonne {j + 1}': row[j]
                 for j in range(len(row))}
        else:
            d = {columns[0]: row[0]}
        rows.append(d)

    # Colonnes numériques (au moins une valeur parseable)
    numeric_cols = []
    for col in columns:
        if any(_to_number(r.get(col)) is not None for r in rows):
            numeric_cols.append(col)

    # Confiance moyenne du rappel (proportion de mots de la question présents)
    q_words = _question_words(question)
    conf = 0.0
    if q_words and rows:
        hits = 0
        for r in rows:
            joined = ' '.join(str(v).lower() for v in r.values())
            hits += len(q_words & set(w for w in re.sub(r'[^\w\s]', ' ', joined).split() if len(w) > 2))
        conf = min(1.0, hits / (len(rows) * max(len(q_words), 1)))

    return {
        'mode': mode,
        'columns': columns,
        'rows': rows,
        'numeric_columns': numeric_cols,
        'count': len(rows),
        'confiance': round(conf, 3),
    }


def _target_column(question: str, numeric_cols: List[str],
                   columns: List[str]) -> Optional[str]:
    """Colonne numérique cible de l'agrégat : nommée dans la question, sinon 1ʳᵉ."""
    if not numeric_cols:
        return None
    q = _column_name(question)
    for hint in _NUM_COLUMN_HINTS:
        if hint in q:
            for col in numeric_cols:
                if hint in _column_name(col):
                    return col
    # Mots de la question dans les noms de colonnes
    q_words = _question_words(question)
    for col in numeric_cols:
        if q_words & set(w for w in re.sub(r'[^\w\s]', ' ', _column_name(col)).split()
                         if len(w) > 2):
            return col
    return numeric_cols[0]


def _aggregates(question: str, analysis: Dict) -> List[Dict]:
    """Calcule les agrégats demandés par la question (compte/somme/moyenne/min/max)."""
    intent = detect_intent(question)
    if not intent:
        return []
    rows, cols = analysis['rows'], analysis['columns']
    numeric_cols = analysis['numeric_columns']

    if intent == 'compte':
        return [{'operation': 'compte', 'libelle': 'Nombre de lignes',
                 'colonne': None, 'valeur': len(rows)}]

    target = _target_column(question, numeric_cols, cols)
    if target is None:
        return [{'operation': intent, 'libelle': f'Agrégat ({intent})',
                 'colonne': None, 'valeur': None}]

    values = [n for n in (_to_number(r.get(target)) for r in rows)
              if n is not None]
    if not values:
        return [{'operation': intent, 'libelle': f'Agrégat ({intent})',
                 'colonne': target, 'valeur': None}]

    ops = {
        'somme': ('Total', sum(values)),
        'moyenne': ('Moyenne', sum(values) / len(values)),
        'max': ('Maximum', max(values)),
        'min': ('Minimum', min(values)),
    }
    lib, val = ops[intent]
    return [{'operation': intent, 'libelle': f'{lib} — {target}',
             'colonne': target, 'valeur': round(val, 2)}]


def query_data(engine, department_id: str, question: str,
               max_rows: int = 500) -> Dict:
    """Répond à une question de DONNÉES : tableau + agrégats (aperçu JSON)."""
    facts = _recall_all(engine, department_id, question, max_facts=max_rows)
    intent = detect_intent(question)
    if intent:
        # 1. Ciblage par le mot le plus DISCRIMINANT de la question (celui qui
        # matche le moins de lignes, avec au moins un match) : « combien de
        # clients actifs ? » → lignes « actif » ; « factures en retard » →
        # lignes « retard ». Sans filtre, l'agrégat porterait sur toute la
        # table (clients + factures + paie mélangées dans un département).
        q_words = _question_words(question)
        if q_words:
            # Compte des matchs par mot : (lignes tabulaires, total) — un mot
            # qui ne matche que du texte libre (bilan, notes) ne discrimine
            # pas une table.
            counts = {}
            for w in q_words:
                tot = sum(1 for f in facts if w in _wordset(f.text))
                tab = sum(1 for f in facts
                          if w in _wordset(f.text) and _is_tabular(f))
                counts[w] = (tab, tot)
            discriminants = [w for w, (tab, tot) in counts.items()
                             if tab > 0 and tot < len(facts)]
            if discriminants:
                best = min(discriminants,
                           key=lambda w: (counts[w][0], counts[w][1]))
                facts = [f for f in facts if best in _wordset(f.text)]
        # 2. L'agrégat porte sur la TABLE (lignes |) quand elle est
        # majoritaire : les faits texte libre (seed, définitions…) ne sont
        # pas des enregistrements comptables.
        tabular = [f for f in facts if _is_tabular(f)]
        if tabular and len(tabular) >= len(facts) / 2:
            facts = tabular
    else:
        # Liste : si le rappel mélange lignes de table et texte libre
        # (ex. « liste des factures en retard » + une phrase contenant
        # « facturés »), la liste de données porte sur la TABLE.
        tabular = [f for f in facts if _is_tabular(f)]
        if tabular and len(tabular) < len(facts):
            facts = tabular
    analysis = analyze_rows(facts, question)
    analysis['aggregates'] = _aggregates(question, analysis)
    analysis['intent'] = intent
    analysis['question'] = question
    analysis['facts_utilises'] = len(facts)
    return analysis


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEL — .xlsx (openpyxl) avec fallback CSV
# ═══════════════════════════════════════════════════════════════════════════════

_GOLD = 'C9A84C'
_DARK = '14141F'


def _slug(text: str) -> str:
    s = _normalize(text)
    s = re.sub(r'[^a-z0-9]+', '_', s).strip('_')
    return s[:40] or 'donnees'


def build_excel(engine, department_id: str, question: str,
                titre: Optional[str] = None) -> Tuple[io.BytesIO, str]:
    """Génère le livrable Excel : feuille Données + feuille Résumé (KPIs)."""
    dept = engine.departments.get(department_id)
    dept_name = dept.name if dept else 'departement'
    data = query_data(engine, department_id, question, max_rows=1000)
    title = titre or f"{dept_name} — {question[:60]}"
    filename = f"{_slug(title)}_{date.today().isoformat()}"

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return _build_csv(data, title, filename)

    wb = openpyxl.Workbook()

    # ── Feuille Données ────────────────────────────────────────────────────────
    ws = wb.active
    ws.title = 'Données'
    ws['A1'] = title
    ws['A1'].font = Font(bold=True, size=13, color=_GOLD)
    ws['A2'] = f"Département : {dept_name} · Généré par KA Enterprise · {date.today().isoformat()}"
    ws['A2'].font = Font(size=9, color='888888')

    header_row = 4
    for j, col in enumerate(data['columns'], start=1):
        c = ws.cell(row=header_row, column=j, value=col)
        c.font = Font(bold=True, color='FFFFFF')
        c.fill = PatternFill('solid', fgColor=_GOLD)
        c.alignment = Alignment(horizontal='center')

    numeric_set = set(data['numeric_columns'])
    for i, row in enumerate(data['rows'], start=header_row + 1):
        for j, col in enumerate(data['columns'], start=1):
            v = row.get(col, '')
            if col in numeric_set and _to_number(v) is not None:
                ws.cell(row=i, column=j, value=_to_number(v))
            else:
                ws.cell(row=i, column=j, value=v if v != '' else None)

    last = header_row + len(data['rows'])
    if data['rows']:
        ws.auto_filter.ref = f"A{header_row}:{get_column_letter(len(data['columns']))}{last}"
    ws.freeze_panes = f"A{header_row + 1}"

    # Largeurs automatiques (cap 42)
    for j, col in enumerate(data['columns'], start=1):
        width = max([len(str(col))] + [len(str(r.get(col, ''))) for r in data['rows']][:200] or [8])
        ws.column_dimensions[get_column_letter(j)].width = min(max(width + 2, 10), 42)

    # ── Feuille Résumé ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet('Résumé')
    ws2['A1'] = title
    ws2['A1'].font = Font(bold=True, size=13, color=_GOLD)
    meta = [
        ('Département', dept_name),
        ('Question', question),
        ('Date', date.today().isoformat()),
        ('Lignes', len(data['rows'])),
        ('Mode de structure', data['mode']),
        ('Colonnes', ', '.join(data['columns'])),
    ]
    for i, (k, v) in enumerate(meta, start=3):
        ws2.cell(row=i, column=1, value=k).font = Font(bold=True, color=_GOLD)
        ws2.cell(row=i, column=2, value=v)

    if data['aggregates']:
        r0 = 3 + len(meta) + 1
        ws2.cell(row=r0, column=1, value='Indicateurs').font = Font(bold=True, size=11)
        for i, agg in enumerate(data['aggregates'], start=r0 + 1):
            ws2.cell(row=i, column=1, value=agg['libelle']).font = Font(bold=True)
            ws2.cell(row=i, column=2, value=agg['valeur'])
    ws2.column_dimensions['A'].width = 30
    ws2.column_dimensions['B'].width = 40

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio, filename + '.xlsx'


def _build_csv(data: Dict, title: str, filename: str) -> Tuple[io.BytesIO, str]:
    """Fallback : CSV (; — locale FR Excel) si openpyxl est indisponible."""
    bio = io.BytesIO()
    w = csv.writer(bio, delimiter=';')
    w.writerow([title])
    w.writerow(data['columns'])
    for row in data['rows']:
        w.writerow([row.get(c, '') for c in data['columns']])
    if data['aggregates']:
        for agg in data['aggregates']:
            w.writerow([agg['libelle'], agg['valeur']])
    bio.seek(0)
    return bio, filename + '.csv'


def export_csv(data: Dict) -> str:
    """Version texte CSV (;) d'une analyse — pour le bouton « .csv »."""
    out = io.StringIO()
    w = csv.writer(out, delimiter=';')
    w.writerow(data['columns'])
    for row in data['rows']:
        w.writerow([row.get(c, '') for c in data['columns']])
    for agg in data.get('aggregates', []):
        w.writerow([agg['libelle'], agg['valeur']])
    return out.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# TEXTES — emails, rapports, comptes-rendus, lettres, notes
# ═══════════════════════════════════════════════════════════════════════════════

def _prose(facts: List) -> str:
    """Faits → prose française connectée (HWAT) puis corrigée (polish_prose)."""
    try:
        from hwat_surface import decode_texts
        lines = decode_texts([f.text for f in facts])
        text = ' '.join(lines) if lines else ' '.join(f.text for f in facts)
    except Exception:
        text = ' '.join(f.text for f in facts)
    try:
        from french_corrector import polish_prose
        text = polish_prose(text)
    except Exception:
        pass
    return text


def _is_tabular(fact) -> bool:
    """Un fait est « tabulaire » si ses segments | sont des cellules (≥2)."""
    return len(_split_segments(fact.text)) >= 2


def _corps_pour(facts: List) -> Tuple[str, str]:
    """
    Corps du document selon la nature des données :
      - faits tabulaires (CSV/tableaux) → liste à puces « • … » ;
      - faits texte libre → prose connectée (HWAT) + correction française.
    Retourne (corps, style).
    """
    if facts and sum(_is_tabular(f) for f in facts) >= max(2, len(facts) // 2):
        return '\n'.join(f'• {f.text}' for f in facts[:20]), 'liste'
    return _prose(facts), 'prose'


def _signature(dept_name: str) -> str:
    return (f"\n\n—\nGénéré par KA Enterprise — IA harmonique\n"
            f"Département : {dept_name}\nCette réponse repose uniquement sur les données privées de votre entreprise.")


def _intro(brief: str, objet: str = '') -> str:
    if objet:
        return f"Suite à votre demande concernant « {brief} », voici les informations relevées :"
    return f"Voici les informations issues de nos données concernant « {brief} » :"


def _closing(facts_count: int) -> str:
    base = "N'hésitez pas à nous solliciter pour tout complément ou une version détaillée."
    return base if facts_count >= 3 else ("Les données disponibles sur ce sujet sont encore limitées — "
                                          "nous vous invitons à les compléter par l'ingestion de documents.")


def _action_facts(facts: List) -> List:
    """Faits orientés « action / échéance » (prochaines étapes d'un CR)."""
    markers = ['prochaine', 'prevu', 'prévu', 'prevoir', 'prévoir', 'echeance',
               'échéance', 'avant le', 'doit', 'action', 'planifie', 'planifié',
               'programme', 'programmé', 'calendrier', 'dans les']
    out = []
    for f in facts:
        low = f.text.lower()
        if any(m in low for m in markers):
            out.append(f)
    return out[:5]


def compose_document(engine, department_id: str, brief: str,
                     doc_format: str = 'rapport',
                     destinataire: Optional[str] = None,
                     objet: Optional[str] = None) -> Dict:
    """
    Prépare un texte structuré (email, rapport, compte-rendu, lettre, note)
    à partir des données privées du département.
    """
    if doc_format not in _FORMATS:
        doc_format = 'rapport'
    dept = engine.departments.get(department_id)
    dept_name = dept.name if dept else 'département'
    facts = _recall_all(engine, department_id, brief, max_facts=25)
    corps, _style = _corps_pour(facts)
    today = date.today().strftime('%d/%m/%Y')
    dest = destinataire or 'Madame, Monsieur'
    obj = objet or f"Compte rendu — {brief[:60]}"
    n = len(facts)

    if doc_format == 'email':
        texte = (f"Objet : {obj}\n\n"
                 f"Bonjour {dest},\n\n"
                 f"{_intro(brief, obj)}\n\n{corps}\n\n"
                 f"{_closing(n)}{_signature(dept_name)}")

    elif doc_format == 'rapport':
        texte = (f"RAPPORT\n{brief[:80].upper()}\n"
                 f"Date : {today} · Département : {dept_name}\n\n"
                 f"1. INTRODUCTION\n{_intro(brief)}\n\n"
                 f"2. ÉLÉMENTS RELEVÉS\n{corps}\n\n"
                 f"3. CONCLUSION\n"
                 f"Sur la base des {n} éléments disponibles, "
                 f"ces informations permettent d'appuyer la décision. {_closing(n)}"
                 f"{_signature(dept_name)}")

    elif doc_format == 'compte_rendu':
        actions = _action_facts(facts)
        actions_txt = ('\n'.join(f'• {f.text}' for f in actions)) if actions else \
            "À définir lors de la prochaine revue."
        texte = (f"COMPTE-RENDU\nObjet : {obj}\n"
                 f"Date : {today} · Département : {dept_name}"
                 f" · Participants : {dest}\n\n"
                 f"POINTS ABORDÉS\n{corps}\n\n"
                 f"DÉCISIONS\nLes éléments ci-dessus constituent la base des décisions "
                 f"retenues par l'équipe.\n\n"
                 f"PROCHAINES ÉTAPES\n{actions_txt}{_signature(dept_name)}")

    elif doc_format == 'lettre':
        texte = (f"Objet : {obj}\nRéférence : {dept_name} / {today}\n\n"
                 f"{dest},\n\n"
                 f"{_intro(brief, obj)}\n\n{corps}\n\n"
                 f"Nous restons à votre disposition pour tout échange complémentaire.\n\n"
                 f"Je vous prie d'agréer, {dest}, l'expression de mes salutations distinguées."
                 f"{_signature(dept_name)}")

    else:  # note
        points = '\n'.join(f'• {f.text}' for f in facts[:10])
        texte = (f"NOTE INTERNE\nObjet : {obj} · Date : {today}\n"
                 f"Département : {dept_name}\n\n"
                 f"{points if points else 'Aucune donnée disponible sur ce sujet.'}"
                 f"{_signature(dept_name)}")

    return {
        'texte': texte,
        'format': doc_format,
        'objet': obj,
        'destinataire': destinataire or '',
        'facts_utilises': n,
        'confiance': round(min(1.0, n / 8.0), 2),
        'department': dept_name,
    }


def document_to_docx(texte: str, filename: str = 'document') -> Tuple[io.BytesIO, str]:
    """Texte structuré → .docx (python-docx), fallback .txt."""
    try:
        import docx as docx_mod
        from docx.shared import Pt
    except ImportError:
        return io.BytesIO(texte.encode('utf-8')), filename + '.txt'

    doc = docx_mod.Document()
    lines = [l for l in texte.split('\n')]
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        upper = stripped.isupper() or (stripped.startswith('RAPPORT') or
                                       stripped.startswith('COMPTE-RENDU') or
                                       stripped.startswith('NOTE INTERNE') or
                                       stripped.startswith('OBJET'))
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        if i == 0 or upper:
            run.bold = True
            run.font.size = Pt(14 if i == 0 else 11)
        if stripped.startswith('• '):
            p.paragraph_format.left_indent = Pt(18)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, filename + '.docx'


# ═══════════════════════════════════════════════════════════════════════════════
# SYNTHÈSE — le savoir d'un département en un texte
# ═══════════════════════════════════════════════════════════════════════════════

def summarize_department(engine, department_id: str,
                         max_facts: int = 30) -> Dict:
    """Synthèse du savoir du département : prose connectée + sources."""
    dept = engine.departments.get(department_id)
    dept_name = dept.name if dept else 'département'
    all_facts = engine.facts.get(department_id, [])
    if not all_facts:
        return {'resume': '', 'facts': 0, 'sources': {}, 'department': dept_name}

    # Faits les plus représentatifs : rappel sur le nom du département, puis
    # élargi aux mots les plus fréquents si besoin.
    facts = _recall_all(engine, department_id, dept_name, max_facts=max_facts)
    if len(facts) < 5:
        counter = Counter()
        for f in all_facts:
            for w in re.sub(r'[^\w\s]', ' ', f.text.lower()).split():
                if len(w) > 4 and w not in _STOPWORDS:
                    counter[w] += 1
        extra = ' '.join(w for w, _ in counter.most_common(12))
        if extra:
            facts = _recall_all(engine, department_id, dept_name + ' ' + extra,
                                max_facts=max_facts)

    corps, _style = _corps_pour(facts)
    sources = Counter(f.source_document for f in all_facts)
    return {
        'resume': f"Synthèse du département « {dept_name} » ({len(all_facts)} faits) :\n\n{corps}",
        'facts': len(all_facts),
        'sources': dict(sources.most_common(10)),
        'department': dept_name,
        'facts_utilises': len(facts),
    }


if __name__ == '__main__':
    # Démonstration / test autonome
    import sys
    from ka_enterprise_core import EnterpriseEngine
    from pathlib import Path
    import tempfile

    tmp = tempfile.mkdtemp()
    eng = EnterpriseEngine(data_dir=str(Path(tmp) / 'data'))
    tenant = eng.create_tenant('Démo', 'demo@entreprise.fr')
    dept = eng.create_department(tenant.id, 'clients')
    corpus = [
        'client_1 | Dupont SA | Paris | 450000',
        'client_2 | Martin & Fils | Lyon | 120500',
        'client_3 | Durand SARL | Marseille | 78000',
        'client_4 | Bernard Conseil | Lille | 234000',
        'client_5 | Petit Distribution | Nantes | 1500',
    ]
    eng.ingest_text(dept.id, '\n'.join(corpus), source='clients.csv')

    for q in ["liste des clients", "combien de clients avons-nous",
              "quel est le chiffre d'affaires total",
              "moyenne des chiffres d'affaires"]:
        d = query_data(eng, dept.id, q)
        print(f"❓ {q} → {d['count']} lignes ({d['mode']}), agrégats : {d['aggregates']}")

    # Mode « clés » (colonnes nommées)
    dept2 = eng.create_department(tenant.id, 'factures')
    eng.ingest_text(dept2.id, '\n'.join([
        'reference : F-2026-001 | montant : 1 200,50 | statut : payee',
        'reference : F-2026-002 | montant : 3 450,00 | statut : en attente',
        'reference : F-2026-003 | montant : 890,25 | statut : payee',
    ]), source='factures.csv')
    d2 = query_data(eng, dept2.id, "liste des factures")
    print(f"\n❓ liste des factures → {d2['count']} lignes ({d2['mode']}), colonnes : {d2['columns']}")
    assert d2['mode'] == 'clefs' and d2['columns'] == ['reference', 'montant', 'statut']
    d3 = query_data(eng, dept2.id, "montant total des factures")
    print(f"❓ montant total → {d3['aggregates']}")
    assert abs(d3['aggregates'][0]['valeur'] - 5540.75) < 0.01

    # Excel : cellules relues
    bio, name = build_excel(eng, dept.id, "liste des clients")
    print(f"\n📊 Excel : {name} ({len(bio.getvalue())} octets)")
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(bio.getvalue()))
    ws = wb['Données']
    assert ws.max_row == 4 + 5, f"attendu 9 lignes, obtenu {ws.max_row}"
    assert ws.cell(row=4, column=1).value == 'Colonne 1'
    assert ws.cell(row=5, column=4).value == 450000  # numérique, pas texte
    print(f"   OK : feuille Données {ws.max_row} lignes, autofilter={ws.auto_filter.ref}")

    # Documents — les 5 formats
    for fmt in ['email', 'rapport', 'compte_rendu', 'lettre', 'note']:
        doc = compose_document(eng, dept.id, "situation des clients", doc_format=fmt)
        print(f"\n✉️ {fmt.upper()} ({doc['facts_utilises']} faits) :\n{doc['texte'][:260]}...")

    # .docx
    doc = compose_document(eng, dept.id, "situation des clients", doc_format='rapport')
    bio2, name2 = document_to_docx(doc['texte'], 'rapport_clients')
    print(f"\n📄 DOCX : {name2} ({len(bio2.getvalue())} octets)")

    # Synthèse
    s = summarize_department(eng, dept.id)
    print(f"\n🧠 SYNTHÈSE ({s['facts']} faits, sources {s['sources']}) :\n{s['resume'][:200]}...")

    print("\n✅ TOUS LES TESTS MODULE PASSENT")
